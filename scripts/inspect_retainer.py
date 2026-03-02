from docx import Document
import os

def inspect_docx(file_path):
    if not os.path.exists(file_path):
        print(f"Error: {file_path} not found")
        return

    doc = Document(file_path)
    print(f"--- Content of {file_path} ---")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    
    # Check tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    print(f"[TABLE] {cell.text}")

if __name__ == "__main__":
    inspect_docx("docs/Richards_Law_Retainer.docx")
